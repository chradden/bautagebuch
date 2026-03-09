"""Word-Dokument (DOCX) – layout-getreu zum PDF-Bericht (tagesbericht.html)."""
import os
import re
from datetime import date, datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config
from core.pdf import _build_total_cost_estimate, _build_eintrag_foto_map

# ── Farben identisch zu tagesbericht.html ─────────────────────────────────

_C_DARK      = "2C3E50"
_C_DARK_RGB  = RGBColor(0x2C, 0x3E, 0x50)
_C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
_C_BODY      = RGBColor(0x33, 0x33, 0x33)
_C_META      = RGBColor(0x66, 0x66, 0x66)
_C_TH_BG     = "EDF2F7"          # Tabellen-Header-Hintergrund
_C_TH_TEXT   = RGBColor(0x1F, 0x2D, 0x3D)
_C_TD_BORDER = "D7DEE7"          # Zellenrahmen
_C_FOTO_BG   = "FCFDFF"          # Bild-/Bildbeschreibungs-Spalte
_C_GREY_TEXT = RGBColor(0x7B, 0x87, 0x94)
_C_DESC_TEXT = RGBColor(0x43, 0x53, 0x64)

_PRIO = {
    "rot": {
        "bg_heading": "FDE2DF",
        "bg_content": "FFF5F5",
        "text":       RGBColor(0x8F, 0x2D, 0x23),
        "border":     "F3C3BE",
    },
    "gelb": {
        "bg_heading": "FFF1BF",
        "bg_content": "FFFBEA",
        "text":       RGBColor(0x7A, 0x5B, 0x00),
        "border":     "EAD68A",
    },
    "gruen": {
        "bg_heading": "DAF2DF",
        "bg_content": "F2FBF5",
        "text":       RGBColor(0x20, 0x60, 0x3A),
        "border":     "B8DDC0",
    },
}

# Spaltenbreiten (A4 Quer, ~27,3 cm nutzbar) = 8% / 32% / 26% / 34%
_COL_WIDTHS = [Cm(2.2), Cm(8.7), Cm(7.1), Cm(9.3)]

# Labels wie im PDF Detail-Block
_DETAIL_LABELS = ["Zustand", "Problem", "Maßnahme", "Dringlichkeit", "Kostenschätzung"]

# ── XML-Hilfsfunktionen ────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = _C_TD_BORDER, size: int = 6) -> None:
    """Setzt Rahmenlinien um eine Tabellenzelle."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), str(size))
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), color)
        tcBorders.append(bd)
    tcPr.append(tcBorders)


def _remove_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "none")
        tcBorders.append(bd)
    tcPr.append(tcBorders)


def _set_para_shading(para, hex_color: str) -> None:
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _set_para_bottom_border(para, hex_color: str = "2C3E50", size: int = 24) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_para_left_border(para, hex_color: str, size: int = 12) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _set_table_no_borders(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "none")
        tblBorders.append(bd)
    tblPr.append(tblBorders)


def _set_para_spacing(para, before: int = 0, after: int = 0) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


def _set_row_cant_split(row) -> None:
    """Verhindert das Aufteilen einer Tabellenzeile über zwei Seiten."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    cantSplit.set(qn("w:val"), "1")
    trPr.append(cantSplit)


# ── Inline-Markup ─────────────────────────────────────────────────────────

def _add_inline_markup(para, text: str, size: float = 10,
                        color: RGBColor | None = None) -> None:
    """Rendert **fett** und *kursiv* Markdown-Markup in einen Paragraphen."""
    for part in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text):
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            run = para.add_run(part)
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color


# ── Tabellenzellen-Renderer ───────────────────────────────────────────────

def _add_detail_cell_content(cell, text: str) -> None:
    """
    Rendert den Details-Block: jedes Label-Wert-Paar in einem eigenen Absatz,
    analog zu den .detail-label Pills im PDF.
    """
    # Markdown-Markup entfernen
    clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    clean = re.sub(r'\*([^*]+)\*', r'\1', clean)

    label_pattern = '|'.join(re.escape(l) for l in _DETAIL_LABELS)
    parts = re.split(rf'({label_pattern})\s*:?\s*', clean)

    if len(parts) <= 1:
        para = cell.paragraphs[0]
        run = para.add_run(clean.strip())
        run.font.size = Pt(9)
        _set_para_spacing(para, before=0, after=20)
        return

    first_para_used = False
    i = 0
    while i < len(parts):
        part = parts[i].strip()
        if not part:
            i += 1
            continue

        if part in _DETAIL_LABELS:
            # Jedes Label bekommt einen eigenen Absatz in der Zelle
            para = cell.paragraphs[0] if not first_para_used else cell.add_paragraph()
            first_para_used = True
            _set_para_spacing(para, before=0, after=10)
            para.paragraph_format.left_indent = Cm(0.15)

            r_label = para.add_run(part)
            r_label.bold = True
            r_label.font.size = Pt(8.5)
            r_label.font.color.rgb = _C_TH_TEXT

            if i + 1 < len(parts):
                value = parts[i + 1].strip().rstrip(" ,;")
                r_sep = para.add_run(": ")
                r_sep.font.size = Pt(9)
                r_sep.font.color.rgb = _C_BODY
                r_val = para.add_run(value)
                r_val.font.size = Pt(9)
                r_val.font.color.rgb = _C_BODY
                i += 2
                continue
        else:
            # Freitext vor erstem Label
            para = cell.paragraphs[0] if not first_para_used else cell.add_paragraph()
            first_para_used = True
            _set_para_spacing(para, before=0, after=10)
            para.paragraph_format.left_indent = Cm(0.15)
            run = para.add_run(part)
            run.font.size = Pt(9)
        i += 1


def _add_bild_cell(cell, photos: list[dict]) -> None:
    """Fügt Fotos in eine DOCX-Tabellenzelle ein (wie Bild-Spalte im PDF)."""
    _set_cell_bg(cell, _C_FOTO_BG)
    if not photos:
        r = cell.paragraphs[0].add_run("Kein Bild")
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = _C_GREY_TEXT
        return

    first = True
    for idx, foto in enumerate(photos, 1):
        path = foto.get("dateipfad_abs", "")
        if not path or not os.path.exists(path):
            continue
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        # Label "Bild 1", "Bild 2" etc.
        r_label = para.add_run(f"Bild {idx}")
        r_label.bold = True
        r_label.font.size = Pt(8)
        r_label.font.color.rgb = _C_DARK_RGB
        _set_para_spacing(para, before=0, after=4)
        # Foto in neuer Zeile – einheitliche Breite 5 cm
        img_para = cell.add_paragraph()
        _set_para_spacing(img_para, before=0, after=20)
        try:
            img_para.add_run().add_picture(path, width=Cm(5.0))
        except Exception:
            img_para.add_run("[Bild nicht verfügbar]").font.size = Pt(8)


def _add_bildbeschreibung_cell(cell, photos: list[dict]) -> None:
    """Fügt Bildbeschreibungen in eine DOCX-Tabellenzelle ein."""
    _set_cell_bg(cell, _C_FOTO_BG)
    if not photos:
        r = cell.paragraphs[0].add_run("Keine Bildbeschreibung")
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = _C_GREY_TEXT
        return

    first = True
    for idx, foto in enumerate(photos, 1):
        desc = (foto.get("beschreibung") or "").strip()
        if not desc:
            continue
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        r_label = para.add_run(f"Bild {idx}")
        r_label.bold = True
        r_label.font.size = Pt(8)
        r_label.font.color.rgb = _C_DARK_RGB
        # Beschreibung auf max. 250 Zeichen kürzen
        short = desc[:250].rstrip()
        if len(desc) > 250:
            short += "…"
        desc_para = cell.add_paragraph()
        r_desc = desc_para.add_run(short)
        r_desc.font.size = Pt(9)
        r_desc.font.color.rgb = _C_DESC_TEXT
        if first:
            first = False


# ── KI-Tabelle als DOCX-Tabelle rendern ───────────────────────────────────

def _add_ki_table(doc: Document, raw_rows: list[str],
                   current_prio: str | None,
                   eintrag_fotos: dict) -> None:
    """
    Wandelt eine Markdown-Tabelle in eine DOCX-Tabelle um –
    identische Spaltenstruktur wie im PDF (Eintrag/Details/Bild/Bildbeschreibung).
    """
    if len(raw_rows) < 2:
        return

    # Header-Zeile parsen
    headers = [c.strip() for c in raw_rows[0].strip("|").split("|") if c.strip()]

    # Alle Folgezeilen einlesen – Separator-Zeilen (|---|---|) überspringen
    data_rows = []
    for row in raw_rows[1:]:
        stripped = row.strip()
        if not stripped:
            continue
        # Separator-Zeile erkennen: enthält nur -, :, |, Leerzeichen
        if re.match(r'^\|[\s:\-|]+\|$', stripped):
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        data_rows.append(cells)

    if not data_rows:
        return

    # Immer 4 Spalten: Eintrag | Details | Bild | Bildbeschreibung
    n_cols = 4
    display_headers = (list(headers) + ["", "", "", ""])[:n_cols]
    display_headers[2] = "Bild"
    display_headers[3] = "Bildbeschreibung"

    colors = _PRIO.get(current_prio) if current_prio else None

    tbl = doc.add_table(rows=0, cols=n_cols)
    # Rahmenstil setzen
    tbl.style = "Table Grid"

    # ── Header-Zeile ──────────────────────────────────────────────────
    hdr_row = tbl.add_row()
    for idx, h in enumerate(display_headers):
        cell = hdr_row.cells[idx]
        bg = colors["bg_heading"] if colors else _C_TH_BG
        _set_cell_bg(cell, bg)
        _set_cell_borders(cell)
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = colors["text"] if colors else _C_TH_TEXT
        para.paragraph_format.left_indent = Cm(0.15)
        _set_para_spacing(para, before=30, after=30)

    # ── Daten-Zeilen ──────────────────────────────────────────────────
    for row_cells in data_rows:
        entry_text = row_cells[0] if row_cells else ""
        nr_match = re.search(r'Nr\.?\s*(\d+)', entry_text)
        photos = []
        if nr_match and eintrag_fotos:
            photos = eintrag_fotos.get(int(nr_match.group(1)), [])

        row = tbl.add_row()
        # Zeile darf nicht über zwei Seiten gehen
        _set_row_cant_split(row)

        for idx in range(n_cols):
            cell = row.cells[idx]
            _set_cell_borders(cell)
            if colors:
                _set_cell_bg(cell, colors["bg_content"])
            para = cell.paragraphs[0]
            para.paragraph_format.left_indent = Cm(0.15)
            _set_para_spacing(para, before=30, after=30)

            src = row_cells[idx] if idx < len(row_cells) else ""

            if idx == 0:
                # Eintrag-Spalte
                run = para.add_run(src)
                run.font.size = Pt(9.5)
                run.bold = True
            elif idx == 1:
                # Details-Spalte: jedes Label in eigenem Absatz
                _add_detail_cell_content(cell, src)
            elif idx == 2:
                # Bild-Spalte – echte Fotos
                _add_bild_cell(cell, photos)
            elif idx == 3:
                # Bildbeschreibung
                _add_bildbeschreibung_cell(cell, photos)

    # ── Spaltenbreiten ────────────────────────────────────────────────
    for row in tbl.rows:
        for idx, width in enumerate(_COL_WIDTHS):
            row.cells[idx].width = width

    # Abstand nach Tabelle
    gap = doc.add_paragraph()
    _set_para_spacing(gap, before=0, after=60)


# ── Header-Block ──────────────────────────────────────────────────────────

def _add_report_header(doc: Document, projekt_name: str, projekt_adresse: str,
                        datum: date, bauleiter_name: str) -> None:
    title_para = doc.add_paragraph()
    run = title_para.add_run("Instandhaltungsplanung")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _C_DARK_RGB
    _set_para_spacing(title_para, before=0, after=0)

    meta_tbl = doc.add_table(rows=1, cols=4)
    _set_table_no_borders(meta_tbl)
    meta_items = [
        ("Objekt",         projekt_name),
        ("Adresse",        projekt_adresse or "–"),
        ("Datum",          datum.strftime("%d.%m.%Y")),
        ("Verantwortlich", bauleiter_name),
    ]
    for idx, (label, value) in enumerate(meta_items):
        cell = meta_tbl.rows[0].cells[idx]
        _remove_cell_borders(cell)
        para = cell.paragraphs[0]
        r_label = para.add_run(label + ": ")
        r_label.bold = True
        r_label.font.size = Pt(9)
        r_label.font.color.rgb = _C_META
        r_value = para.add_run(value)
        r_value.font.size = Pt(9)
        r_value.font.color.rgb = _C_META

    sep = doc.add_paragraph()
    _set_para_bottom_border(sep, hex_color=_C_DARK, size=24)
    _set_para_spacing(sep, before=60, after=120)


# ── Abschnitts-Header (weißer Text auf dunkelblauem Hintergrund) ──────────

def _add_section_header(doc: Document, text: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, _C_DARK)
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = _C_WHITE
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Cm(0.3)
    gap = doc.add_paragraph()
    _set_para_spacing(gap, before=0, after=60)


# ── Prio-Erkennung ────────────────────────────────────────────────────────

def _detect_prio(text: str) -> str | None:
    lower = text.lower()
    if "rot" in lower or "🔴" in text or "sofort" in lower:
        return "rot"
    if "gelb" in lower or "🟡" in text or "zeitnah" in lower:
        return "gelb"
    if "grün" in lower or "gruen" in lower or "🟢" in text or "geplant" in lower:
        return "gruen"
    return None


# ── KI-Bericht Renderer ───────────────────────────────────────────────────

def _render_ki_bericht(doc: Document, markdown_text: str,
                        eintrag_fotos: dict | None = None) -> None:
    """
    Wandelt den KI-Markdown-Bericht in DOCX um.
    Markdown-Tabellen werden als echte DOCX-Tabellen mit Foto-Injektion gerendert.
    Prio-Abschnitte erhalten farbige Hintergründe wie im PDF.
    """
    if eintrag_fotos is None:
        eintrag_fotos = {}

    current_prio: str | None = None
    lines = markdown_text.splitlines()
    i = 0

    while i < len(lines):
        stripped = lines[i].strip()

        # ── Leerzeile ──────────────────────────────────────────────
        if not stripped:
            gap = doc.add_paragraph()
            _set_para_spacing(gap, before=0, after=40)
            i += 1
            continue

        # ── Markdown-Tabelle erkennen (aktuelle + nächste Zeile ---|---|) ──
        if stripped.startswith("|") and i + 1 < len(lines):
            next_stripped = lines[i + 1].strip()
            if re.match(r'^\|[-| :]+\|', next_stripped):
                # Tabelle sammeln
                table_rows = [stripped]
                i += 1
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_rows.append(lines[i].strip())
                    i += 1
                _add_ki_table(doc, table_rows, current_prio, eintrag_fotos)
                continue

        # ── H1 ─────────────────────────────────────────────────────
        m = re.match(r'^# (.*)', stripped)
        if m:
            current_prio = _detect_prio(m.group(1)) or current_prio
            para = doc.add_paragraph()
            run = para.add_run(m.group(1))
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = _C_DARK_RGB
            _set_para_spacing(para, before=120, after=60)
            i += 1
            continue

        # ── H2 ─────────────────────────────────────────────────────
        m = re.match(r'^## (.*)', stripped)
        if m:
            text = m.group(1)
            current_prio = _detect_prio(text)
            colors = _PRIO.get(current_prio) if current_prio else None
            para = doc.add_paragraph()
            if colors:
                _set_para_shading(para, colors["bg_heading"])
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = colors["text"] if colors else _C_DARK_RGB
            para.paragraph_format.left_indent = Cm(0.3)
            _set_para_spacing(para, before=120, after=30)
            i += 1
            continue

        # ── H3 ─────────────────────────────────────────────────────
        m = re.match(r'^### (.*)', stripped)
        if m:
            text = m.group(1)
            prio = _detect_prio(text)
            if prio:
                current_prio = prio
            colors = _PRIO.get(current_prio) if current_prio else None
            para = doc.add_paragraph()
            if colors:
                _set_para_shading(para, colors["bg_heading"])
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = colors["text"] if colors else _C_DARK_RGB
            para.paragraph_format.left_indent = Cm(0.3)
            _set_para_spacing(para, before=80, after=20)
            i += 1
            continue

        # ── Aufzählungspunkt ────────────────────────────────────────
        m = re.match(r'^[-*] (.*)', stripped)
        if m:
            colors = _PRIO.get(current_prio) if current_prio else None
            para = doc.add_paragraph()
            if colors:
                _set_para_shading(para, colors["bg_content"])
                _set_para_left_border(para, colors["border"])
            para.paragraph_format.left_indent = Cm(0.8)
            para.paragraph_format.first_line_indent = Cm(-0.35)
            bullet_run = para.add_run("• ")
            bullet_run.font.size = Pt(10)
            if colors:
                bullet_run.font.color.rgb = colors["text"]
            _add_inline_markup(para, m.group(1), size=10)
            _set_para_spacing(para, before=20, after=20)
            i += 1
            continue

        # ── Normaler Absatz ─────────────────────────────────────────
        colors = _PRIO.get(current_prio) if current_prio else None
        para = doc.add_paragraph()
        if colors:
            _set_para_shading(para, colors["bg_content"])
            _set_para_left_border(para, colors["border"])
            para.paragraph_format.left_indent = Cm(0.3)
        _add_inline_markup(para, stripped, size=10)
        _set_para_spacing(para, before=20, after=20)
        i += 1


# ── Gesamtkosten-Box ──────────────────────────────────────────────────────

def _add_kosten_box(doc: Document, gesamt: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    _set_table_no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, "F8FAFC")
    _remove_cell_borders(cell)

    label_para = cell.paragraphs[0]
    r_label = label_para.add_run("Gesamtkostenschätzung")
    r_label.bold = True
    r_label.font.size = Pt(9)
    r_label.font.color.rgb = _C_DARK_RGB

    val_para = cell.add_paragraph()
    r_val = val_para.add_run(gesamt)
    r_val.bold = True
    r_val.font.size = Pt(12)
    r_val.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)

    label_para.paragraph_format.left_indent = Cm(0.3)
    val_para.paragraph_format.left_indent = Cm(0.3)
    _set_para_spacing(label_para, before=60, after=20)
    _set_para_spacing(val_para, before=0, after=60)


# ── Foto-Dokumentation ────────────────────────────────────────────────────

def _add_foto_section(doc: Document, fotos: list) -> None:
    foto_liste = [
        f for f in fotos
        if getattr(f, "dateipfad", None) and os.path.exists(os.path.abspath(f.dateipfad))
    ]
    if not foto_liste:
        return

    doc.add_paragraph()
    _add_section_header(doc, "Fotodokumentation")

    foto_tbl = doc.add_table(rows=0, cols=3)
    _set_table_no_borders(foto_tbl)

    row_cells = None
    for idx, foto in enumerate(foto_liste):
        col = idx % 3
        if col == 0:
            row_cells = foto_tbl.add_row().cells
            for c in row_cells:
                _remove_cell_borders(c)

        cell = row_cells[col]
        try:
            cell.paragraphs[0].add_run().add_picture(
                os.path.abspath(foto.dateipfad), width=Cm(8.5)
            )
        except Exception:
            cell.paragraphs[0].add_run("[Bild nicht verfügbar]")

        parts = []
        uhrzeit = getattr(foto, "uhrzeit", None)
        if uhrzeit:
            try:
                parts.append(uhrzeit.strftime("%H:%M"))
            except AttributeError:
                parts.append(str(uhrzeit))
        beschreibung = (getattr(foto, "beschreibung", "") or "")
        if beschreibung:
            parts.append(beschreibung[:120])
        if parts:
            cap = cell.add_paragraph(" – ".join(parts))
            if cap.runs:
                cap.runs[0].font.size = Pt(8)
                cap.runs[0].italic = True
                cap.runs[0].font.color.rgb = _C_META


# ── Footer ────────────────────────────────────────────────────────────────

def _add_footer(doc: Document) -> None:
    sep = doc.add_paragraph()
    _set_para_bottom_border(sep, hex_color="CCCCCC", size=4)
    _set_para_spacing(sep, before=200, after=40)

    tbl = doc.add_table(rows=1, cols=2)
    _set_table_no_borders(tbl)
    for cell in tbl.rows[0].cells:
        _remove_cell_borders(cell)

    left_para = tbl.rows[0].cells[0].paragraphs[0]
    r_left = left_para.add_run(f"Erstellt am {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    r_left.font.size = Pt(8)
    r_left.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    right_para = tbl.rows[0].cells[1].paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_right = right_para.add_run("Instandhaltungsplanung – automatisch generiert")
    r_right.font.size = Pt(8)
    r_right.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ── Hauptfunktion ─────────────────────────────────────────────────────────

def generiere_docx(
    projekt_name: str,
    bauleiter_name: str,
    datum: date,
    eintraege_text: list,
    fotos: list,
    ki_bericht: str = "",
    projekt_adresse: str = "",
) -> str:
    """
    Generiert einen editierbaren Word-Tagesbericht (.docx) im PDF-Layout.
    Markdown-Tabellen des KI-Berichts werden als DOCX-Tabellen mit echten
    Fotos gerendert – identisch zur PDF-Darstellung.
    """
    doc = Document()

    # A4 Querformat + Ränder 1.2 cm (identisch zum PDF)
    for section in doc.sections:
        section.page_width    = Cm(29.7)
        section.page_height   = Cm(21.0)
        section.top_margin    = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin   = Cm(1.2)
        section.right_margin  = Cm(1.2)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].font.color.rgb = _C_BODY

    # Foto-Map aufbauen (Eintragsnummer → Fotos)
    eintrag_fotos = _build_eintrag_foto_map(eintraege_text)

    # 1. Header
    _add_report_header(doc, projekt_name, projekt_adresse, datum, bauleiter_name)

    # 2. KI-Analyse mit Tabellen und Fotos
    if ki_bericht:
        _add_section_header(doc, "KI-Analyse & Priorisierung")
        _render_ki_bericht(doc, ki_bericht, eintrag_fotos)

        gesamt = _build_total_cost_estimate(eintraege_text)
        if gesamt:
            doc.add_paragraph()
            _add_kosten_box(doc, gesamt)

    # 3. Footer
    _add_footer(doc)

    # Speichern
    os.makedirs(config.PDF_OUTPUT_DIR, exist_ok=True)
    safe_name = projekt_name.replace(" ", "_").replace("/", "-")
    dateiname = f"Tagesbericht_{safe_name}_{datum.strftime('%Y-%m-%d')}.docx"
    docx_pfad = os.path.join(config.PDF_OUTPUT_DIR, dateiname)
    doc.save(docx_pfad)
    return docx_pfad
